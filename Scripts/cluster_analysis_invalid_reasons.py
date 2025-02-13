import pandas as pd
import matplotlib.pyplot as plt
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.manifold import TSNE
from sklearn.cluster import KMeans
from collections import defaultdict
from docx import Document
import re

# Load data
df = pd.read_csv('invalid-applications-sample-reasons.xlsx - invalid-application-sample-reas.csv')

# Define themes for clustering with more granular categories for Missing Documents
themes = {
    "Incorrect Fee": r"(fee|payment|underpayment|overpayment)",
    "Missing Plans": r"(site plan|floor plan|elevation)",
    "Missing Reports": r"(report|statement|assessment|survey)",
    "Missing Forms": r"(form|certificate|ownership)",
    "Validation Checklist": r"(checklist|validation|requirement)",
    "Missing Details": r"(details|clarify|information)",
    "Missing Drawings": r"(drawing|design|sketch|diagram)",
    "Other": r".*"
}

# Initialise a dictionary to hold the themes and their corresponding rows
grouped_reasons = defaultdict(list)

# Iterate over the rows and classify based on themes
for index, row in df.iterrows():
    reason = row["Invalid Reason Details"]
    if pd.isna(reason):
        continue
    matched = False
    for theme, pattern in themes.items():
        if re.search(pattern, reason, re.IGNORECASE):
            grouped_reasons[theme].append(reason)
            matched = True
            break
    if not matched:
        grouped_reasons["Other"].append(reason)

# Separate "Incorrect Fee" into "Incorrect Fee - Underpayment" if specific words are present
incorrect_fee = grouped_reasons.get("Incorrect Fee", [])
underpayment_fee = [reason for reason in incorrect_fee if re.search(r"insufficient|further fee", reason, re.IGNORECASE)]
remaining_fee = [reason for reason in incorrect_fee if reason not in underpayment_fee]
grouped_reasons["Incorrect Fee - Underpayment"] = underpayment_fee
grouped_reasons["Incorrect Fee - Other"] = remaining_fee

# Remove the old "Incorrect Fee" key
del grouped_reasons["Incorrect Fee"]

# Reorder groups so "Incorrect Fee - Other" and "Other" appear at the end
ordered_keys = [key for key in grouped_reasons if key not in ["Incorrect Fee - Other", "Other"]]
ordered_keys += ["Incorrect Fee - Other", "Other"]

# Create a Word document
doc = Document()
doc.add_heading('Grouped Invalid Reason Details', level=1)

# Add each group and its count to the document
for theme in ordered_keys:
    reasons = grouped_reasons[theme]
    count = len(reasons)
    if theme == "Incorrect Fee - Other":
        underpayment_count = len(grouped_reasons.get("Incorrect Fee - Underpayment", []))
        doc.add_heading(f"{theme} ({count} instances, Underpayment: {underpayment_count} instances):", level=2)
    else:
        doc.add_heading(f"{theme} ({count} instances):", level=2)

# Save the document
doc.save('Grouped_Invalid_Reason_Details.docx')

# Perform TF-IDF vectorisation
all_reasons = df["Invalid Reason Details"].dropna().tolist()
vectorizer = TfidfVectorizer(stop_words='english', max_features=500)
tfidf_matrix = vectorizer.fit_transform(all_reasons)

# Reduce dimensions using TSNE
tsne = TSNE(n_components=2, random_state=42, perplexity=30, n_iter=1000)
tsne_results = tsne.fit_transform(tfidf_matrix.toarray())

# Cluster the data into 10 groups using KMeans
kmeans = KMeans(n_clusters=10, random_state=42)
clusters = kmeans.fit_predict(tfidf_matrix)

# Prepare data for visualization
cluster_colors = clusters

plt.figure(figsize=(10, 8))
scatter = plt.scatter(tsne_results[:, 0], tsne_results[:, 1], c=cluster_colors, cmap='tab10', alpha=0.7)
plt.colorbar(scatter, ticks=range(10), label='Clusters')
plt.title('TSNE Visualization of Clusters')
plt.xlabel('TSNE Dimension 1')
plt.ylabel('TSNE Dimension 2')
plt.savefig('TSNE_Clusters.png')
plt.show()

# Create a CSV file with themes as columns and reasons as rows
max_rows = max(len(reasons) for reasons in grouped_reasons.values())
output_data = {theme: grouped_reasons[theme] + [None] * (max_rows - len(grouped_reasons[theme])) for theme in ordered_keys}
output_df = pd.DataFrame(output_data)
output_df.to_csv('Grouped_Invalid_Reason_Details.csv', index=False)

print("Outputs generated:")
print("1. Document 'Grouped_Invalid_Reason_Details.docx' created with grouped reasons.")
print("2. Visualisation 'TSNE_Clusters.png' saved.")
print("3. CSV file 'Grouped_Invalid_Reason_Details.csv' created with grouped reasons.")
