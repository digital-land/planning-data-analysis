import os
import geopandas as gpd
import pandas as pd
import folium
import asyncio
from pyppeteer import launch
from docx import Document
from docx.shared import Inches, Pt

async def process_dataset(dataset):
    """
    Processes all shapefiles in a dataset folder by analysing each, generating individual reports,
    and then merging them into a single consolidated DOCX report.

    Args:
        dataset (str): Name of the dataset folder.
    """
    
    dataset_folder = f"datasets/{dataset}"
    folder_path = f"{dataset_folder}/files"

    def find_geospatial_files(folder_path):
        """Finds geospatial files (.shp, .gpkg, .csv) in the folder."""
        return [file for file in os.listdir(folder_path) if file.endswith(('.shp', '.gpkg', '.csv'))]

    print(f"Looking for geospatial files in: {folder_path}")
    geospatial_files = find_geospatial_files(folder_path)

    if not geospatial_files:
        print("No geospatial files found.")
        return

    print("Geospatial files found:", geospatial_files)

    report_files = []

    def analyse_folder(folder_path):
        """Analyses the folder and returns file sizes."""
        details = {file: f"Size: {os.path.getsize(os.path.join(folder_path, file))} bytes"
                   for file in os.listdir(folder_path)}
        details["Total Folder Size"] = f"{sum(os.path.getsize(os.path.join(folder_path, file)) for file in os.listdir(folder_path))} bytes"
        return details

    folder_description = analyse_folder(folder_path)

    for index, geospatial_file in enumerate(geospatial_files):
        print(f"\nProcessing: {geospatial_file} ({index + 1}/{len(geospatial_files)})")

        geojson_output_file = f"{dataset_folder}/{dataset}_{geospatial_file}_sample.geojson"
        data_plot_file = f"{dataset_folder}/{dataset}_{geospatial_file}_plot.png"
        report_output_file = f"{dataset_folder}/{dataset}_{geospatial_file}_analysis_report.docx"

        def load_geospatial_file(folder_path, file_name):
            """Loads a geospatial file (.shp, .gpkg, .csv) into a GeoDataFrame."""
            file_path = os.path.join(folder_path, file_name)
            try:
                if file_name.endswith(('.shp', '.gpkg')):
                    return gpd.read_file(file_path, engine="pyogrio", on_invalid="ignore")
                elif file_name.endswith('.csv'):
                    return pd.read_csv(file_path)
                else:
                    raise ValueError("Unsupported file format.")
            except Exception as e:
                print(f"Error loading file: {e}")
                return None

        gdf = load_geospatial_file(folder_path, geospatial_file)
        if gdf is None:
            continue

        def analyse_geospatial_file(gdf, geojson_output_file):
            """Analyses geospatial data and saves a sample as GeoJSON."""
            try:
                gdf.head(50).to_file(geojson_output_file, driver="GeoJSON")
                print(f"Sample GeoJSON saved: {geojson_output_file}")
            except Exception as e:
                print(f"Error saving GeoJSON: {e}")
                return None

            return {
                "geometry_type": ', '.join(gdf.geom_type.unique()),
                "record_count": len(gdf),
                "features": ', '.join(gdf.columns),
                "crs": str(gdf.crs) if gdf.crs else "Unknown",
                "file_size": os.path.getsize(os.path.join(folder_path, geospatial_file)),
                "geojson_file": geojson_output_file
            }

        geospatial_data = analyse_geospatial_file(gdf, geojson_output_file)
        if geospatial_data is None:
            continue

        CHROME_PATH = "C:/Program Files/Google/Chrome/Application/chrome.exe"

        m = gdf.iloc[:1].explore()
        map_html_path = f"{dataset_folder}/map_{geospatial_file}.html"
        m.save(map_html_path)

        async def capture_map():
            """Captures a screenshot of an interactive map using a headless browser."""
            try:
                browser = await launch(headless=True, executablePath=CHROME_PATH, args=["--no-sandbox"])
                page = await browser.newPage()
                await page.goto(f"file://{os.path.abspath(map_html_path)}")
                await asyncio.sleep(3)
                await page.screenshot({"path": data_plot_file, "fullPage": True})
                print(f"Map saved: {data_plot_file}")
            except Exception as e:
                print(f"Error capturing map: {e}")
            finally:
                await browser.close()

        await capture_map()

        def generate_report(folder_data, geospatial_data, output_file, image_file, first_report):
            """
            Generates a report summarising the dataset, including metadata, folder contents, 
            a map plot, and a GeoJSON sample.

            Args:
                folder_data (dict): Folder contents.
                geospatial_data (dict): Geospatial metadata.
                output_file (str): Report file path.
                image_file (str): Dataset plot path.
                first_report (bool): Whether this is the first report (keeps the main heading & folder contents).
            """
            doc = Document()

            if first_report:
                doc.add_heading(f"{dataset} Analysis Report", 0)
                
                # Folder contents only in the first report
                doc.add_heading('Folder Contents:', level=1)
                for file, details in folder_data.items():
                    doc.add_paragraph(f"{file}: {details}", style='List Bullet')

            doc.add_heading(f"{geospatial_file} Analysis:", level=1)
            doc.add_paragraph(f"Geometry Type: {geospatial_data['geometry_type']}", style='List Bullet')
            doc.add_paragraph(f"Number of Records: {geospatial_data['record_count']}", style='List Bullet')
            doc.add_paragraph(f"Features: {geospatial_data['features']}", style='List Bullet')
            doc.add_paragraph(f"CRS: {geospatial_data['crs']}", style='List Bullet')
            doc.add_paragraph(f"File Size: {geospatial_data['file_size']} bytes", style='List Bullet')

            doc.add_heading("Dataset Plot", level=1)
            if os.path.exists(image_file):
                doc.add_picture(image_file, width=Inches(6))
                doc.add_paragraph(f"Figure: First-row dataset visualisation ({image_file})")
            else:
                doc.add_paragraph("No image available for visualisation.")

            doc.add_heading("GeoJSON Sample Data", level=1)
            if os.path.exists(geospatial_data["geojson_file"]):
                with open(geospatial_data["geojson_file"], 'r') as geojson_file:
                    doc.add_paragraph("Sample GeoJSON content:")
                    doc.add_paragraph(geojson_file.read(2000))
            else:
                doc.add_paragraph("No GeoJSON file available.")

            doc.save(output_file)
            print(f"Report saved: {output_file}")
            return doc

        first_report = (index == 0)
        generate_report(folder_description, geospatial_data, report_output_file, data_plot_file, first_report)
        report_files.append(report_output_file)

    if len(report_files) > 1:
        final_report_path = f"{dataset_folder}/{dataset}_final_analysis_report.docx"
        merged_doc = Document(report_files[0])

        for report in report_files[1:]:
            doc_to_merge = Document(report)
            merged_doc.add_page_break()
            doc_heading = f"{geospatial_files[report_files.index(report)]} Analysis:"
            merged_doc.add_heading(doc_heading, level=1)
            for element in doc_to_merge.paragraphs:
                merged_doc.add_paragraph(element.text)

        merged_doc.save(final_report_path)
        print(f"\nFinal merged report saved: {final_report_path}")

# Example usage:

#dataset = ""
#await process_dataset(dataset)
