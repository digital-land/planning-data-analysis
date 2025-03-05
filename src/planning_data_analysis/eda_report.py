import asyncio
import os

import geopandas as gpd
import pandas as pd
from docx import Document
from docx.shared import Inches
from playwright.async_api import async_playwright


def find_geospatial_files(folder_path):
    """Finds geospatial files (.shp, .gpkg, .csv) in the folder."""
    return [
        file
        for file in os.listdir(folder_path)
        if file.endswith((".shp", ".gpkg", ".csv"))
    ]


def analyse_folder(folder_path):
    """Analyses the folder and returns file sizes."""
    details = {
        file: f"Size: {os.path.getsize(os.path.join(folder_path, file))} bytes"
        for file in os.listdir(folder_path)
    }
    details[
        "Total Folder Size"
    ] = f"{sum(os.path.getsize(os.path.join(folder_path, file)) for file in os.listdir(folder_path))} bytes"
    return details


def load_geospatial_file(folder_path, file_name):
    """Loads a geospatial file (.shp, .gpkg, .csv) into a GeoDataFrame."""
    file_path = os.path.join(folder_path, file_name)
    try:
        if file_name.endswith((".shp", ".gpkg")):
            return gpd.read_file(file_path, engine="pyogrio", on_invalid="ignore")
        elif file_name.endswith(".csv"):
            return pd.read_csv(file_path)
        else:
            raise ValueError("Unsupported file format.")
    except Exception as e:
        print(f"Error loading file: {e}")
        return None


def analyse_geospatial_file(gdf, geojson_output_file, folder_path, geospatial_file):
    """
    Analyses geospatial data and saves a sample as GeoJSON.

    Parameters:
    ----------
    gdf : GeoDataFrame
        The geospatial data to analyze
    geojson_output_file : str
        Path to save the GeoJSON sample
    folder_path : str
        Path to the folder containing the original file
    geospatial_file : str
        Name of the geospatial file being analyzed
    """
    try:
        gdf.head(50).to_file(geojson_output_file, driver="GeoJSON")
        print(f"Sample GeoJSON saved: {geojson_output_file}")
    except Exception as e:
        print(f"Error saving GeoJSON: {e}")
        return None

    return {
        "geometry_type": ", ".join(gdf.geom_type.unique()),
        "record_count": len(gdf),
        "features": ", ".join(gdf.columns),
        "crs": str(gdf.crs) if gdf.crs else "Unknown",
        "file_size": os.path.getsize(os.path.join(folder_path, geospatial_file)),
        "geojson_file": geojson_output_file,
    }


def generate_report(
    folder_data,
    geospatial_data,
    output_file,
    image_file,
    first_report,
    dataset,
    geospatial_file,
):
    """
    Generates a report summarising the dataset, including metadata, folder contents,
    a map plot, and a GeoJSON sample.

    Args:
        folder_data (dict): Folder contents.
        geospatial_data (dict): Geospatial metadata.
        output_file (str): Report file path.
        image_file (str): Dataset plot path.
        first_report (bool): Whether this is the first report (keeps the main heading & folder contents).
        dataset (str): Name of the dataset.
        geospatial_file (str): Name of the geospatial file being analyzed.
    """
    doc = Document()

    if first_report:
        doc.add_heading(f"{dataset} Analysis Report", 0)

        # Folder contents only in the first report
        doc.add_heading("Folder Contents:", level=1)
        for file, details in folder_data.items():
            doc.add_paragraph(f"{file}: {details}", style="List Bullet")

    doc.add_heading(f"{geospatial_file} Analysis:", level=1)
    doc.add_paragraph(
        f"Geometry Type: {geospatial_data['geometry_type']}", style="List Bullet"
    )
    doc.add_paragraph(
        f"Number of Records: {geospatial_data['record_count']}", style="List Bullet"
    )
    doc.add_paragraph(f"Features: {geospatial_data['features']}", style="List Bullet")
    doc.add_paragraph(f"CRS: {geospatial_data['crs']}", style="List Bullet")
    doc.add_paragraph(
        f"File Size: {geospatial_data['file_size']} bytes", style="List Bullet"
    )

    doc.add_heading("Dataset Plot", level=1)
    if os.path.exists(image_file):
        doc.add_picture(image_file, width=Inches(6))
        doc.add_paragraph(f"Figure: First-row dataset visualisation ({image_file})")
    else:
        doc.add_paragraph("No image available for visualisation.")

    doc.add_heading("GeoJSON Sample Data", level=1)
    if os.path.exists(geospatial_data["geojson_file"]):
        with open(geospatial_data["geojson_file"], "r") as geojson_file:
            doc.add_paragraph("Sample GeoJSON content:")
            doc.add_paragraph(geojson_file.read(2000))
    else:
        doc.add_paragraph("No GeoJSON file available.")

    doc.save(output_file)
    print(f"Report saved: {output_file}")
    return doc


async def process_dataset(dataset, input_dir, output_dir):
    """
    Processes all shapefiles in a dataset folder by analysing each, generating individual reports,
    and then merging them into a single consolidated DOCX report.

    Args:
        dataset (str): Name of the dataset folder.
        input_dir (str): Path to input directory containing geospatial files.
        output_dir (str): Path to output directory for reports and visualizations.
    """
    # Create output directory if it doesn't exist
    os.makedirs(output_dir, exist_ok=True)

    print(f"Looking for geospatial files in: {input_dir}")
    geospatial_files = find_geospatial_files(input_dir)

    if not geospatial_files:
        print("No geospatial files found.")
        return

    print("Geospatial files found:", geospatial_files)

    report_files = []
    folder_description = analyse_folder(input_dir)

    for index, geospatial_file in enumerate(geospatial_files):
        print(f"\nProcessing: {geospatial_file} ({index + 1}/{len(geospatial_files)})")

        geojson_output_file = os.path.join(
            output_dir, f"{dataset}_{geospatial_file}_sample.geojson"
        )
        data_plot_file = os.path.join(
            output_dir, f"{dataset}_{geospatial_file}_plot.png"
        )
        report_output_file = os.path.join(
            output_dir, f"{dataset}_{geospatial_file}_analysis_report.docx"
        )

        gdf = load_geospatial_file(input_dir, geospatial_file)
        if gdf is None:
            continue

        geospatial_data = analyse_geospatial_file(
            gdf, geojson_output_file, input_dir, geospatial_file
        )
        if geospatial_data is None:
            continue

        m = gdf.iloc[:1].explore()
        map_html_path = os.path.join(output_dir, f"map_{geospatial_file}.html")
        m.save(map_html_path)

        async with async_playwright() as p:
            browser = await p.chromium.launch()
            page = await browser.new_page()
            await page.goto(f"file://{os.path.abspath(map_html_path)}")
            await page.wait_for_timeout(3000)  # Wait for 3 seconds
            await page.screenshot(path=data_plot_file, full_page=True)
            print(f"Map saved: {data_plot_file}")
            await browser.close()

        first_report = index == 0
        generate_report(
            folder_description,
            geospatial_data,
            report_output_file,
            data_plot_file,
            first_report,
            dataset,
            geospatial_file,
        )
        report_files.append(report_output_file)

    if len(report_files) > 1:
        final_report_path = os.path.join(
            output_dir, f"{dataset}_final_analysis_report.docx"
        )
        merged_doc = Document(report_files[0])

        for report in report_files[1:]:
            doc_to_merge = Document(report)
            merged_doc.add_page_break()
            doc_heading = f"{geospatial_files[report_files.index(report)]} Analysis:"
            merged_doc.add_heading(doc_heading, level=1)
            for element in doc_to_merge.paragraphs:
                merged_doc.add_paragraph(element.text)

        merged_doc.save(final_report_path)
        print(f"\nFinal merged report saved: {final_report_path}")


def generate_eda_report(dataset, input_dir, output_dir):
    """
    Generate an exploratory data analysis report for geospatial data.

    Parameters:
    ----------
    dataset : str
        Name of the dataset
    input_dir : str
        Path to input directory containing geospatial files
    output_dir : str
        Path to output directory for reports and visualizations
    """
    asyncio.run(process_dataset(dataset, input_dir, output_dir))
